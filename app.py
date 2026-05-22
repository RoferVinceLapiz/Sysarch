from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, send_file, Response
import sqlite3
import os
from datetime import date, datetime
from io import BytesIO, StringIO
import csv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from openpyxl import Workbook
from flask import request, jsonify

lab_maintenance = {}


def normalize_lab(lab):
    if lab is None:
        return ''
    return str(lab).replace('Lab', '').strip()


def is_pc_under_maintenance(lab, pc_number):
    lab = normalize_lab(lab)

    try:
        pc_number = int(pc_number)
    except (TypeError, ValueError):
        return False

    maintenance_pcs = lab_maintenance.get(lab, [])

    return pc_number in [int(pc) for pc in maintenance_pcs]


def block_if_pc_maintenance(lab, pc_number):
    return is_pc_under_maintenance(lab, pc_number)

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'

DB_PATH = 'database.db'
UPLOAD_FOLDER = 'static/uploads'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# =============================================================
# DOCX CUSTOM VISUAL FORMATTING HELPERS
# =============================================================

def apply_cell_background(cell, hex_color):
    """Applies a modern background fill color to an individual table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shading)

def apply_cell_padding(cell, top=100, bottom=100, left=150, right=150):
    """Optional: Adds clean internal spacing/padding inside a table cell (in twentieths of a point)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def docx_cell_background(cell, hex_color):
    """Sets background shading tint for individual cell containers."""
    shading_xml = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_xml)

def docx_cell_padding(cell, top=100, bottom=100, left=120, right=120):
    """Configures cell buffer padding heights/widths in dxa units."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_side, spacing_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_side}')
        node.set(qn('w:w'), str(spacing_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def docx_clean_borders(table):
    """Replaces default box lines with subtle, modern light-gray row dividers."""
    tblPr = table._tbl.tblPr
    borders_str = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'  <w:insideV w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    borders_xml = parse_xml(borders_str)
    tblPr.append(borders_xml)


# =============================================================
# HELPERS
# =============================================================

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def get_sitin_report_rows(conn):
    return conn.execute("""
        SELECT
            sr.id,
            sr.id_number,
            sr.purpose,
            sr.lab,
            sr.login_time,
            sr.logout_time,
            sr.evaluation_cleanliness,
            sr.evaluation_hours,
            sr.evaluation_task,
            sr.evaluation_points,
            sr.evaluation_note,
            s.first_name,
            s.last_name,
            s.course,
            s.course_level,
            CASE
                WHEN sr.logout_time IS NOT NULL THEN
                    CAST((julianday(sr.logout_time) - julianday(sr.login_time)) * 24 * 60 AS INTEGER)
                ELSE
                    NULL
            END AS duration_minutes
        FROM sitin_records sr
        JOIN students s ON sr.id_number = s.id_number
        ORDER BY sr.login_time DESC
    """).fetchall()
    
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def get_local_time():
    return datetime.now().strftime('%Y-%m-%d %H:%M:%S')

def get_last_session_reset(conn):
    reset = conn.execute("""
        SELECT reset_at
        FROM session_resets
        ORDER BY reset_at DESC
        LIMIT 1
    """).fetchone()

    return reset['reset_at'] if reset else None


def get_student_session_info(conn, id_number):
    BASE_SESSIONS = 30  # Base sessions after reset

    last_reset = get_last_session_reset(conn)

    if last_reset:
        # Count sit-ins after last reset
        completed_sitins = conn.execute("""
            SELECT COUNT(*) AS count
            FROM sitin_records
            WHERE id_number = ?
              AND logout_time IS NOT NULL
              AND login_time >= ?
        """, (id_number, last_reset)).fetchone()['count']

        raw_points = conn.execute("""
            SELECT COALESCE(SUM(evaluation_points), 0) AS points
            FROM sitin_records
            WHERE id_number = ?
              AND logout_time IS NOT NULL
              AND login_time >= ?
        """, (id_number, last_reset)).fetchone()['points']
    else:
        completed_sitins = conn.execute("""
            SELECT COUNT(*) AS count
            FROM sitin_records
            WHERE id_number = ?
              AND logout_time IS NOT NULL
        """, (id_number,)).fetchone()['count']

        raw_points = conn.execute("""
            SELECT COALESCE(SUM(evaluation_points), 0) AS points
            FROM sitin_records
            WHERE id_number = ?
              AND logout_time IS NOT NULL
        """, (id_number,)).fetchone()['points']

    raw_points = float(raw_points or 0)
    bonus_sessions = int(raw_points // 3)

    # Remaining sessions after reset
    remaining_sessions = BASE_SESSIONS - completed_sitins + bonus_sessions
    if remaining_sessions < 0:
        remaining_sessions = 0

    return {
        'completed_sitins': completed_sitins,
        'raw_points': raw_points,
        'awarded_points': raw_points,
        'bonus_sessions': bonus_sessions,
        'remaining_sessions': remaining_sessions
    }

def add_notification(id_number, message):
    if not id_number or not message:
        return

    conn = get_db()
    conn.execute('''
        INSERT INTO notifications (id_number, message)
        VALUES (?, ?)
    ''', (id_number, message))
    conn.commit()
    conn.close()


def notify_all_students(message):
    if not message:
        return

    conn = get_db()

    students = conn.execute('''
        SELECT id_number
        FROM students
    ''').fetchall()

    for student in students:
        conn.execute('''
            INSERT INTO notifications (id_number, message)
            VALUES (?, ?)
        ''', (student['id_number'], message))

    conn.commit()
    conn.close()


def get_student_notifications(id_number):
    conn = get_db()

    notifications = conn.execute('''
        SELECT *
        FROM notifications
        WHERE id_number = ?
        ORDER BY created_at DESC
        LIMIT 10
    ''', (id_number,)).fetchall()

    notification_count = conn.execute('''
        SELECT COUNT(*) AS count
        FROM notifications
        WHERE id_number = ?
          AND is_read = 0
    ''', (id_number,)).fetchone()['count']

    conn.close()

    return notifications, notification_count


def is_reservation_enabled(conn):
    """Check if reservations are currently enabled for students."""
    setting = conn.execute('''
        SELECT setting_value
        FROM settings
        WHERE setting_key = ?
    ''', ('reservations_enabled',)).fetchone()
    
    return setting and setting['setting_value'] == '1' if setting else True


def toggle_reservation_status(conn):
    """Toggle the reservation status between enabled and disabled."""
    current_status = conn.execute('''
        SELECT setting_value
        FROM settings
        WHERE setting_key = ?
    ''', ('reservations_enabled',)).fetchone()
    
    new_status = '0' if (current_status and current_status['setting_value'] == '1') else '1'
    
    conn.execute('''
        UPDATE settings
        SET setting_value = ?, updated_at = CURRENT_TIMESTAMP
        WHERE setting_key = ?
    ''', (new_status, 'reservations_enabled'))
    
    conn.commit()
    return new_status == '1'


# =============================================================
# DATABASE
# =============================================================

def init_db():
    conn = get_db()
    cursor = conn.cursor()

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            id_number TEXT UNIQUE NOT NULL,
            last_name TEXT NOT NULL,
            first_name TEXT NOT NULL,
            middle_name TEXT,
            course TEXT NOT NULL,
            course_level TEXT NOT NULL,
            password TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            address TEXT
        )
    ''')

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS session_resets (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            reset_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    """)

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS admins (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS announcements (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            message TEXT NOT NULL,
            posted_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sitin_records (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            id_number TEXT NOT NULL,
            purpose TEXT NOT NULL,
            lab TEXT NOT NULL,
            login_time DATETIME DEFAULT CURRENT_TIMESTAMP,
            logout_time DATETIME,
            rating INTEGER,
            feedback TEXT
        )
    ''')

    cursor.execute("PRAGMA table_info(sitin_records)")
    sitin_columns = [column[1] for column in cursor.fetchall()]

    if 'evaluation_cleanliness' not in sitin_columns:
        cursor.execute('ALTER TABLE sitin_records ADD COLUMN evaluation_cleanliness INTEGER DEFAULT 0')

    if 'evaluation_hours' not in sitin_columns:
        cursor.execute('ALTER TABLE sitin_records ADD COLUMN evaluation_hours INTEGER DEFAULT 0')

    if 'evaluation_task' not in sitin_columns:
        cursor.execute('ALTER TABLE sitin_records ADD COLUMN evaluation_task INTEGER DEFAULT 0')

    if 'evaluation_points' not in sitin_columns:
        cursor.execute('ALTER TABLE sitin_records ADD COLUMN evaluation_points REAL DEFAULT 0')

    if 'evaluation_note' not in sitin_columns:
        cursor.execute('ALTER TABLE sitin_records ADD COLUMN evaluation_note TEXT')

    if 'pc_number' not in sitin_columns:
        cursor.execute('ALTER TABLE sitin_records ADD COLUMN pc_number INTEGER')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS notifications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            id_number TEXT NOT NULL,
            message TEXT NOT NULL,
            is_read INTEGER DEFAULT 0,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS lab_applications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            lab TEXT NOT NULL,
            application_name TEXT NOT NULL,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(lab, application_name)
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS reservations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            id_number TEXT NOT NULL,
            purpose TEXT NOT NULL,
            lab TEXT NOT NULL,
            pc_number INTEGER,
            time_in TIME NOT NULL,
            date DATE NOT NULL,
            status TEXT DEFAULT 'Pending',
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    cursor.execute("PRAGMA table_info(reservations)")
    reservation_columns = [column[1] for column in cursor.fetchall()]

    if 'pc_number' not in reservation_columns:
        cursor.execute('ALTER TABLE reservations ADD COLUMN pc_number INTEGER')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS settings (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            setting_key TEXT UNIQUE NOT NULL,
            setting_value TEXT NOT NULL,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # Initialize default reservation status if not exists
    cursor.execute('''
        INSERT OR IGNORE INTO settings (setting_key, setting_value)
        VALUES (?, ?)
    ''', ('reservations_enabled', '1'))

    cursor.execute('''
        INSERT OR IGNORE INTO admins (username, password)
        VALUES (?, ?)
    ''', ('admin', 'admin123'))

    conn.commit()
    conn.close()
# =============================================================
# STUDENT ROUTES
# =============================================================

@app.route('/', methods=['GET', 'POST'])
def login():
    conn = get_db()

    leaderboard_rows = conn.execute("""
        SELECT
            s.id_number,
            s.first_name,
            s.last_name,
            s.course,
            s.course_level,
            COUNT(sr.id) AS completed_sitins,
            COALESCE(SUM(sr.evaluation_points), 0) AS raw_points
        FROM students s
        LEFT JOIN sitin_records sr
            ON s.id_number = sr.id_number
           AND sr.logout_time IS NOT NULL
        GROUP BY
            s.id_number,
            s.first_name,
            s.last_name,
            s.course,
            s.course_level
        ORDER BY
            raw_points DESC,
            completed_sitins DESC,
            s.last_name ASC,
            s.first_name ASC
        LIMIT 3
    """).fetchall()

    if request.method == 'POST':
        id_number = request.form['id_number']
        password = request.form['password']

        admin = conn.execute(
            'SELECT * FROM admins WHERE username = ? AND password = ?',
            (id_number, password)
        ).fetchone()

        if admin:
            session['admin_id'] = admin['id']
            session['admin_user'] = admin['username']
            conn.close()
            flash('Welcome back, Administrator!', 'success')
            return redirect(url_for('admin_dashboard'))

        student = conn.execute(
            'SELECT * FROM students WHERE id_number = ? AND password = ?',
            (id_number, password)
        ).fetchone()

        if student:
            session['student_id'] = student['id_number']
            session['student_name'] = student['first_name']

            conn.close()
            flash(f"Welcome! {student['first_name']} {student['last_name']}", 'success')
            return redirect(url_for('dashboard'))

        conn.close()
        flash('Invalid ID number or password.', 'error')
        return render_template('login.html', top_students=leaderboard_rows)

    conn.close()
    return render_template('login.html', top_students=leaderboard_rows)


@app.route('/check-pc-availability')
def check_pc_availability():
    if 'student_id' not in session:
        return jsonify({'unavailable_pcs': []})

    lab = request.args.get('lab')
    date_value = request.args.get('date')
    time_in = request.args.get('time_in')

    if not lab or not date_value or not time_in:
        return jsonify({'unavailable_pcs': []})

    conn = get_db()

    # Select PCs that are either:
    # 1. Currently occupied by a sit-in (logout_time IS NULL)
    # 2. Already reserved for that lab/date/time (Pending or Approved)
    reserved_pcs = conn.execute("""
        SELECT pc_number
        FROM reservations
        WHERE lab = ?
          AND date = ?
          AND time_in = ?
          AND status IN ('Pending', 'Approved')
          AND pc_number IS NOT NULL
        UNION
        SELECT pc_number
        FROM sitin_records
        WHERE lab = ?
          AND pc_number IS NOT NULL
          AND logout_time IS NULL
    """, (lab, date_value, time_in, lab)).fetchall()

    conn.close()

    unavailable_pcs = [row['pc_number'] for row in reserved_pcs if row['pc_number'] is not None]

    maintenance_pcs = lab_maintenance.get(str(lab), [])
    unavailable_pcs = list(set(unavailable_pcs + maintenance_pcs))

    return jsonify({'unavailable_pcs': unavailable_pcs})


@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        id_number = request.form['id_number']
        last_name = request.form['last_name']
        first_name = request.form['first_name']
        middle_name = request.form['middle_name']
        course = request.form['course']
        course_level = request.form['course_level']
        password = request.form['password']
        repeat_password = request.form['repeat_password']
        email = request.form['email']
        address = request.form['address']

        if password != repeat_password:
            flash('Passwords do not match.', 'error')
            return redirect(url_for('register'))

        try:
            conn = get_db()
            conn.execute('''
                INSERT INTO students
                (id_number, last_name, first_name, middle_name, course, course_level, password, email, address)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (id_number, last_name, first_name, middle_name, course, course_level, password, email, address))
            conn.commit()
            conn.close()
            flash('Registration successful! Please login.', 'success')
            return redirect(url_for('login'))

        except sqlite3.IntegrityError:
            flash('ID Number or Email already exists.', 'error')
            return redirect(url_for('register'))

    return render_template('register.html')


@app.route('/dashboard')
def dashboard():
    if 'student_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    student = conn.execute(
        'SELECT * FROM students WHERE id_number = ?',
        (session['student_id'],)
    ).fetchone()

    session_info = get_student_session_info(conn, session['student_id'])
    remaining = session_info['remaining_sessions']

    announcements = conn.execute(
        'SELECT * FROM announcements ORDER BY posted_at DESC'
    ).fetchall()

    last_reset = get_last_session_reset(conn)

    if last_reset:
        leaderboard_rows = conn.execute('''
            SELECT
                s.id_number,
                s.first_name,
                s.last_name,
                s.course,
                s.course_level,

                COUNT(sr.id) AS completed_sitins,

                COALESCE(SUM(sr.evaluation_points), 0) AS raw_points,

                CAST(COALESCE(SUM(sr.evaluation_points), 0) / 3 AS INTEGER) AS bonus_sessions

            FROM students s
            LEFT JOIN sitin_records sr
                ON s.id_number = sr.id_number
               AND sr.logout_time IS NOT NULL
               AND sr.login_time >= ?

            GROUP BY
                s.id_number,
                s.first_name,
                s.last_name,
                s.course,
                s.course_level

            ORDER BY
                raw_points DESC,
                completed_sitins DESC,
                s.last_name ASC,
                s.first_name ASC
        ''', (last_reset,)).fetchall()
    else:
        leaderboard_rows = conn.execute('''
            SELECT
                s.id_number,
                s.first_name,
                s.last_name,
                s.course,
                s.course_level,

                COUNT(sr.id) AS completed_sitins,

                COALESCE(SUM(sr.evaluation_points), 0) AS raw_points,

                CAST(COALESCE(SUM(sr.evaluation_points), 0) / 3 AS INTEGER) AS bonus_sessions

            FROM students s
            LEFT JOIN sitin_records sr
                ON s.id_number = sr.id_number
               AND sr.logout_time IS NOT NULL

            GROUP BY
                s.id_number,
                s.first_name,
                s.last_name,
                s.course,
                s.course_level

            ORDER BY
                raw_points DESC,
                completed_sitins DESC,
                s.last_name ASC,
                s.first_name ASC
        ''').fetchall()

    leaderboard = []
    leaderboard_rank = None
    current_leaderboard_info = None
    previous_score = None
    displayed_rank = 0

    for index, row in enumerate(leaderboard_rows, start=1):
        raw_points = float(row['raw_points'] or 0)
        completed_sitins = int(row['completed_sitins'] or 0)
        bonus_sessions = int(row['bonus_sessions'] or 0)

        if previous_score is None or raw_points != previous_score:
            displayed_rank = index

        student_rank = {
            'rank': displayed_rank,
            'id_number': row['id_number'],
            'first_name': row['first_name'],
            'last_name': row['last_name'],
            'course': row['course'],
            'course_level': row['course_level'],
            'completed_sitins': completed_sitins,
            'raw_points': raw_points,
            'awarded_points': raw_points,
            'bonus_sessions': bonus_sessions
        }

        leaderboard.append(student_rank)

        if row['id_number'] == session['student_id']:
            leaderboard_rank = displayed_rank
            current_leaderboard_info = student_rank

        previous_score = raw_points

    top_leaderboard = leaderboard[:10]
    total_ranked_students = len(leaderboard)

    conn.close()

    notifications, notification_count = get_student_notifications(session['student_id'])

    return render_template(
        'dashboard.html',
        student=student,
        announcements=announcements,
        remaining=remaining,
        bonus_sessions=session_info['bonus_sessions'],
        completed_sitins=session_info['completed_sitins'],
        raw_points=session_info['raw_points'],
        awarded_points=session_info['raw_points'],
        leaderboard_rank=leaderboard_rank,
        current_leaderboard_info=current_leaderboard_info,
        top_leaderboard=top_leaderboard,
        total_ranked_students=total_ranked_students,
        notifications=notifications,
        notification_count=notification_count
    )

@app.route('/edit-profile', methods=['GET', 'POST'])
def edit_profile():
    if 'student_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    student = conn.execute(
        'SELECT * FROM students WHERE id_number = ?',
        (session['student_id'],)
    ).fetchone()

    photo_path = f"uploads/{session['student_id']}.png"
    has_photo = os.path.exists(os.path.join('static', photo_path))

    if request.method == 'POST':
        last_name = request.form['last_name']
        first_name = request.form['first_name']
        middle_name = request.form['middle_name']
        course_level = request.form['course_level']
        email = request.form['email']
        course = request.form['course']
        address = request.form['address']

        conn.execute('''
            UPDATE students SET
                last_name = ?, first_name = ?, middle_name = ?,
                course_level = ?, email = ?, course = ?, address = ?
            WHERE id_number = ?
        ''', (
            last_name,
            first_name,
            middle_name,
            course_level,
            email,
            course,
            address,
            session['student_id']
        ))

        conn.commit()
        conn.close()

        add_notification(
            session['student_id'],
            'Your profile information was updated successfully.'
        )

        flash('Profile updated successfully!', 'success')
        return redirect(url_for('edit_profile'))

    conn.close()

    notifications, notification_count = get_student_notifications(session['student_id'])

    return render_template(
        'edit_profile.html',
        student=student,
        photo_path=photo_path,
        has_photo=has_photo,
        notifications=notifications,
        notification_count=notification_count
    )


@app.route('/upload-photo', methods=['POST'])
def upload_photo():
    if 'student_id' not in session:
        return redirect(url_for('login'))

    file = request.files.get('photo')

    if file and allowed_file(file.filename):
        filename = f"{session['student_id']}.png"
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))

        add_notification(
            session['student_id'],
            'Your profile photo was updated successfully.'
        )

        flash('Profile photo updated successfully!', 'success')
    else:
        flash('Invalid file. Please upload a JPG, PNG, or GIF.', 'error')

    return redirect(url_for('edit_profile'))


@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))


@app.route('/api/leaderboard')
def api_leaderboard():
    conn = get_db()

    last_reset = get_last_session_reset(conn)

    if last_reset:
        rows = conn.execute('''
            SELECT
                s.id_number,
                s.first_name,
                s.last_name,
                s.course,
                s.course_level,
                s.email,
                s.address,
                COUNT(sr.id) AS completed_sitins,
                COALESCE(SUM(sr.evaluation_points), 0) AS raw_points
            FROM students s
            LEFT JOIN sitin_records sr
                ON s.id_number = sr.id_number
               AND sr.logout_time IS NOT NULL
               AND sr.login_time >= ?
            GROUP BY s.id_number
            ORDER BY raw_points DESC, completed_sitins DESC, s.last_name ASC, s.first_name ASC
            LIMIT 50
        ''', (last_reset,)).fetchall()
    else:
        rows = conn.execute('''
            SELECT
                s.id_number,
                s.first_name,
                s.last_name,
                s.course,
                s.course_level,
                s.email,
                s.address,
                COUNT(sr.id) AS completed_sitins,
                COALESCE(SUM(sr.evaluation_points), 0) AS raw_points
            FROM students s
            LEFT JOIN sitin_records sr
                ON s.id_number = sr.id_number
               AND sr.logout_time IS NOT NULL
            GROUP BY s.id_number
            ORDER BY raw_points DESC, completed_sitins DESC, s.last_name ASC, s.first_name ASC
            LIMIT 50
        ''').fetchall()

    leaderboard = []
    previous_score = None
    displayed_rank = 0

    for index, row in enumerate(rows, start=1):
        raw_points = float(row['raw_points'] or 0)
        completed_sitins = int(row['completed_sitins'] or 0)

        if previous_score is None or raw_points != previous_score:
            displayed_rank = index

        profile_picture_url = None
        for ext in ('png', 'jpg', 'jpeg', 'webp'):
            profile_pic_path = os.path.join(app.root_path, 'static', 'uploads', f"{row['id_number']}.{ext}")
            if os.path.exists(profile_pic_path):
                profile_picture_url = f"/static/uploads/{row['id_number']}.{ext}"
                break

        leaderboard.append({
            'rank': displayed_rank,
            'id_number': row['id_number'],
            'first_name': row['first_name'],
            'last_name': row['last_name'],
            'course': row['course'],
            'course_level': row['course_level'],
            'email': row['email'],
            'address': row['address'],
            'completed_sitins': completed_sitins,
            'raw_points': raw_points,
            'profile_picture': profile_picture_url
        })

        previous_score = raw_points

    conn.close()

    return jsonify({'leaderboard': leaderboard[:3]})


@app.route('/students')
def students():
    if 'student_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()
    students = conn.execute('SELECT * FROM students').fetchall()
    conn.close()

    notifications, notification_count = get_student_notifications(session['student_id'])

    return render_template(
        'students.html',
        students=students,
        notifications=notifications,
        notification_count=notification_count
    )


# =============================================================
# RESERVATION ROUTES
# =============================================================

@app.route('/reservation', methods=['GET', 'POST'])
def reservation():
    if 'student_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    student = conn.execute(
        'SELECT * FROM students WHERE id_number = ?',
        (session['student_id'],)
    ).fetchone()

    session_info = get_student_session_info(conn, session['student_id'])
    remaining = session_info['remaining_sessions']

    labs = ['524', '526', '528', '530', '542', '544']

    lab_application_rows = conn.execute("""
        SELECT id, lab, application_name
        FROM lab_applications
        ORDER BY lab ASC, application_name ASC
    """).fetchall()

    lab_applications = {lab: [] for lab in labs}

    for row in lab_application_rows:
        lab_key = str(row['lab'])

        if lab_key not in lab_applications:
            lab_applications[lab_key] = []

        lab_applications[lab_key].append({
            'id': row['id'],
            'lab': lab_key,
            'application_name': row['application_name']
        })

    if request.method == 'POST':
        # Check if reservations are enabled
        if not is_reservation_enabled(conn):
            flash('Reservations are currently disabled. Please try again later.', 'error')
            conn.close()
            return redirect(url_for('reservation'))

        purpose = request.form.get('purpose', '').strip()
        lab = request.form.get('lab', '').strip()
        time_in = request.form.get('time_in', '').strip()
        reservation_date = request.form.get('date', '').strip()
        pc_number = request.form.get('pc_number', '').strip()

        if not purpose or not lab or not time_in or not reservation_date or not pc_number:
            flash('Please complete all reservation steps, including PC number.', 'error')
            conn.close()
            return redirect(url_for('reservation'))

        # Block past date and past time
        try:
            selected_datetime = datetime.strptime(
                f"{reservation_date} {time_in}",
                "%Y-%m-%d %H:%M"
            )

            if selected_datetime < datetime.now():
                flash('You cannot reserve using a past date or past time.', 'error')
                conn.close()
                return redirect(url_for('reservation'))

        except ValueError:
            flash('Invalid reservation date or time.', 'error')
            conn.close()
            return redirect(url_for('reservation'))

        if lab not in labs:
            flash('Invalid laboratory selected.', 'error')
            conn.close()
            return redirect(url_for('reservation'))

        selected_application = conn.execute("""
            SELECT id
            FROM lab_applications
            WHERE lab = ?
              AND application_name = ?
        """, (lab, purpose)).fetchone()

        if not selected_application:
            flash(f'{purpose} is not available in Lab {lab}. Please choose an available application.', 'error')
            conn.close()
            return redirect(url_for('reservation'))

        try:
            pc_number_int = int(pc_number)
        except ValueError:
            flash('Invalid PC number selected.', 'error')
            conn.close()
            return redirect(url_for('reservation'))

        if pc_number_int < 1 or pc_number_int > 50:
            flash('PC number must be between 1 and 50.', 'error')
            conn.close()
            return redirect(url_for('reservation'))

        if is_pc_under_maintenance(lab, pc_number_int):
            flash(f'PC {pc_number_int} is currently under maintenance.', 'error')
            conn.close()
            return redirect(url_for('reservation'))

        if remaining <= 0:
            flash('You have no remaining sessions to reserve.', 'error')
            conn.close()
            return redirect(url_for('reservation'))

        existing_pc = conn.execute('''
            SELECT id
            FROM reservations
            WHERE lab = ?
              AND date = ?
              AND time_in = ?
              AND pc_number = ?
              AND status IN ('Pending', 'Approved')
        ''', (lab, reservation_date, time_in, pc_number_int)).fetchone()

        if existing_pc:
            flash(f'PC {pc_number_int} is already reserved for that lab, date, and time.', 'error')
            conn.close()
            return redirect(url_for('reservation'))

        conn.execute('''
            INSERT INTO reservations (id_number, purpose, lab, pc_number, time_in, date, status)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (
            session['student_id'],
            purpose,
            lab,
            pc_number_int,
            time_in,
            reservation_date,
            'Pending'
        ))

        conn.commit()
        conn.close()

        add_notification(
            session['student_id'],
            f'Your reservation was submitted and is pending admin approval. Purpose: {purpose}, Lab: {lab}, PC: {pc_number_int}, Date: {reservation_date}, Time: {time_in}.'
        )

        flash('Reservation submitted successfully! Please wait for admin approval.', 'success')
        return redirect(url_for('reservation'))

    reservations = conn.execute(
        'SELECT * FROM reservations WHERE id_number = ? ORDER BY created_at DESC',
        (session['student_id'],)
    ).fetchall()

    reservations_enabled = is_reservation_enabled(conn)

    conn.close()

    notifications, notification_count = get_student_notifications(session['student_id'])

    return render_template(
        'reservation.html',
        student=student,
        remaining=remaining,
        completed_sitins=session_info['completed_sitins'],
        bonus_sessions=session_info['bonus_sessions'],
        reservations=reservations,
        lab_applications=lab_applications,
        reservations_enabled=reservations_enabled,
        notifications=notifications,
        notification_count=notification_count
    )

@app.route('/reservation/pc-status')
def reservation_pc_status():
    if 'student_id' not in session:
        return jsonify({"success": False, "message": "Unauthorized"}), 401

    lab = request.args.get('lab', '').strip()
    reservation_date = request.args.get('date', '').strip()
    time_in = request.args.get('time_in', '').strip()

    if not lab or not reservation_date or not time_in:
        return jsonify({
            "success": True,
            "unavailable": [],
            "maintenance": [],
            "reserved": []
        })

    conn = get_db()

    reservation_rows = conn.execute('''
        SELECT pc_number
        FROM reservations
        WHERE lab = ?
          AND date = ?
          AND time_in = ?
          AND pc_number IS NOT NULL
          AND status IN ('Pending', 'Approved')
    ''', (lab, reservation_date, time_in)).fetchall()

    sitin_rows = conn.execute('''
        SELECT pc_number
        FROM sitin_records
        WHERE lab = ?
          AND logout_time IS NULL
          AND pc_number IS NOT NULL
    ''', (lab,)).fetchall()

    conn.close()

    reserved_set = set()

    for row in reservation_rows:
        if row['pc_number'] is not None:
            reserved_set.add(int(row['pc_number']))

    for row in sitin_rows:
        if row['pc_number'] is not None:
            reserved_set.add(int(row['pc_number']))

    maintenance_set = set()

    for pc in lab_maintenance.get(str(lab), []):
        maintenance_set.add(int(pc))

    unavailable_set = reserved_set.union(maintenance_set)

    return jsonify({
        "success": True,
        "unavailable": sorted(list(unavailable_set)),
        "maintenance": sorted(list(maintenance_set)),
        "reserved": sorted(list(reserved_set))
    })
    
@app.route('/sit-summary')
def sit_summary():
    if 'student_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    last_reset = get_last_session_reset(conn)

    if last_reset:
        records = conn.execute('''
            SELECT 
                id,
                id_number,
                purpose,
                lab,
                login_time,
                logout_time,
                CAST(
                    (julianday(logout_time) - julianday(login_time)) * 24 * 60
                    AS INTEGER
                ) AS duration_minutes
            FROM sitin_records
            WHERE id_number = ?
              AND logout_time IS NOT NULL
              AND login_time >= ?
            ORDER BY login_time DESC
        ''', (session['student_id'], last_reset)).fetchall()
    else:
        records = conn.execute('''
            SELECT 
                id,
                id_number,
                purpose,
                lab,
                login_time,
                logout_time,
                CAST(
                    (julianday(logout_time) - julianday(login_time)) * 24 * 60
                    AS INTEGER
                ) AS duration_minutes
            FROM sitin_records
            WHERE id_number = ?
              AND logout_time IS NOT NULL
            ORDER BY login_time DESC
        ''', (session['student_id'],)).fetchall()

    total_minutes = 0
    total_sessions = len(records)
    longest_minutes = 0
    sessions = []

    for r in records:
        minutes = r['duration_minutes'] or 0
        total_minutes += minutes

        if minutes > longest_minutes:
            longest_minutes = minutes

        sessions.append({
            'date': r['login_time'][:10],
            'time_in': r['login_time'][11:16],
            'time_out': r['logout_time'][11:16],
            'duration': f"{minutes // 60} hr {minutes % 60} min",
            'purpose': r['purpose'],
            'lab': r['lab']
        })

    average_minutes = total_minutes // total_sessions if total_sessions > 0 else 0
    session_info = get_student_session_info(conn, session['student_id'])

    conn.close()

    notifications, notification_count = get_student_notifications(session['student_id'])

    return render_template(
        'sit_summary.html',
        total_sitin_hours=f"{total_minutes // 60} hrs {total_minutes % 60} min",
        total_sessions=total_sessions,
        average_duration=f"{average_minutes // 60} hr {average_minutes % 60} min",
        longest_session=f"{longest_minutes // 60} hr {longest_minutes % 60} min",
        sessions=sessions,
        remaining=session_info['remaining_sessions'],
        bonus_sessions=session_info['bonus_sessions'],
        notifications=notifications,
        notification_count=notification_count
    )


# =============================================================
# ADMIN ROUTES
# =============================================================

@app.route('/admin/dashboard')
def admin_dashboard():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    students = conn.execute('SELECT * FROM students').fetchall()
    total_students = len(students)

    announcements = conn.execute(
        'SELECT * FROM announcements ORDER BY posted_at DESC'
    ).fetchall()

    course_stats = [
        {"course": row["course"], "count": row["count"]}
        for row in conn.execute(
            'SELECT course, COUNT(*) as count FROM students GROUP BY course'
        ).fetchall()
    ]

    total_sitin = conn.execute(
        'SELECT COUNT(*) as c FROM sitin_records'
    ).fetchone()['c']

    current_sitin = conn.execute(
        'SELECT COUNT(*) as c FROM sitin_records WHERE logout_time IS NULL'
    ).fetchone()['c']

    conn.close()

    return render_template(
        'admin_dashboard.html',
        students=students,
        total_students=total_students,
        announcements=announcements,
        course_stats=course_stats,
        total_sitin=total_sitin,
        current_sitin=current_sitin,
        admin_user=session['admin_user']
    )


from flask import request, jsonify

@app.route('/admin/save-maintenance', methods=['POST'])
def save_maintenance():
    data = request.json
    lab = data['lab']
    pcs = data['pcs']

    # save to DB or memory
    save_to_database(lab, pcs)

    return jsonify({"status": "success"})

@app.route('/admin/lab-maintenance/status')
def lab_maintenance_status():
    if 'admin_id' not in session and 'student_id' not in session:
        return jsonify({}), 401

    return jsonify(lab_maintenance)

@app.route('/admin/add-student', methods=['POST'])
def admin_add_student():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    id_number = request.form.get('id_number', '').strip()
    last_name = request.form.get('last_name', '').strip()
    first_name = request.form.get('first_name', '').strip()
    middle_name = request.form.get('middle_name', '').strip()
    course = request.form.get('course', '').strip()
    course_level = request.form.get('course_level', '').strip()
    password = request.form.get('password', '').strip()
    email = request.form.get('email', '').strip()
    address = request.form.get('address', '').strip()

    if not id_number or not last_name or not first_name or not course or not course_level or not password or not email:
        flash('Please complete all required student fields.', 'error')
        return redirect(url_for('admin_students'))

    try:
        conn = get_db()
        conn.execute('''
            INSERT INTO students
            (id_number, last_name, first_name, middle_name, course, course_level, password, email, address)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            id_number,
            last_name,
            first_name,
            middle_name,
            course,
            course_level,
            password,
            email,
            address
        ))

        conn.commit()
        conn.close()

        add_notification(
            id_number,
            'Your student account has been created by the admin. You can now log in and use the sit-in system.'
        )

        flash(f'Student {first_name} {last_name} added successfully.', 'success')

    except sqlite3.IntegrityError:
        flash('ID Number or Email already exists.', 'error')

    return redirect(url_for('admin_students'))


@app.route('/admin/sitin-report')
def admin_sitin_report():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    total_sitin = conn.execute("""
        SELECT COUNT(*) AS count
        FROM sitin_records
    """).fetchone()['count']

    completed_sitin = conn.execute("""
        SELECT COUNT(*) AS count
        FROM sitin_records
        WHERE logout_time IS NOT NULL
    """).fetchone()['count']

    active_sitin = conn.execute("""
        SELECT COUNT(*) AS count
        FROM sitin_records
        WHERE logout_time IS NULL
    """).fetchone()['count']

    total_raw_points = conn.execute("""
        SELECT COALESCE(SUM(evaluation_points), 0) AS points
        FROM sitin_records
        WHERE logout_time IS NOT NULL
    """).fetchone()['points']

    purpose_stats = conn.execute("""
        SELECT purpose, COUNT(*) AS count
        FROM sitin_records
        GROUP BY purpose
        ORDER BY count DESC
    """).fetchall()

    lab_stats = conn.execute("""
        SELECT lab, COUNT(*) AS count
        FROM sitin_records
        GROUP BY lab
        ORDER BY lab ASC
    """).fetchall()

    top_students = conn.execute("""
        SELECT
            s.id_number,
            s.first_name,
            s.last_name,
            s.course,
            COUNT(sr.id) AS completed_sitins,
            COALESCE(SUM(sr.evaluation_points), 0) AS raw_points,
            CAST(COALESCE(SUM(sr.evaluation_points), 0) / 3 AS INTEGER) AS extra_sessions
        FROM students s
        LEFT JOIN sitin_records sr
            ON s.id_number = sr.id_number
           AND sr.logout_time IS NOT NULL
        GROUP BY s.id_number, s.first_name, s.last_name, s.course
        ORDER BY raw_points DESC, completed_sitins DESC
        LIMIT 10
    """).fetchall()

    report_records = get_sitin_report_rows(conn)

    # Students for Quick Select Search Modal
    student_rows = conn.execute("""
        SELECT
            id_number,
            first_name,
            last_name,
            course,
            course_level
        FROM students
        ORDER BY last_name ASC, first_name ASC
    """).fetchall()

    students = []

    for student in student_rows:
        session_info = get_student_session_info(conn, student['id_number'])

        profile_picture_filename = f"{student['id_number']}.png"
        profile_picture_path = os.path.join(app.config['UPLOAD_FOLDER'], profile_picture_filename)

        students.append({
            'id_number': student['id_number'],
            'first_name': student['first_name'],
            'last_name': student['last_name'],
            'course': student['course'],
            'course_level': student['course_level'],
            'remaining_sessions': session_info['remaining_sessions'],
            'completed_sitins': session_info['completed_sitins'],
            'raw_points': session_info['raw_points'],
            'bonus_sessions': session_info['bonus_sessions'],
            'profile_picture': profile_picture_filename if os.path.exists(profile_picture_path) else None
        })

    conn.close()

    return render_template(
        'admin_sitin_report.html',
        total_sitin=total_sitin,
        completed_sitin=completed_sitin,
        active_sitin=active_sitin,
        total_raw_points=total_raw_points,
        purpose_stats=purpose_stats,
        lab_stats=lab_stats,
        top_students=top_students,
        report_records=report_records,
        students=students,
        admin_user=session['admin_user']
    )

@app.route('/admin/lab-maintenance/save', methods=['POST'])
def save_lab_maintenance():
    if 'admin_id' not in session:
        return jsonify({
            "status": "error",
            "message": "Unauthorized"
        }), 401

    data = request.get_json() or {}

    lab = str(data.get('lab', '')).replace('Lab', '').strip()
    pcs = data.get('pcs', [])

    clean_pcs = []

    for pc in pcs:
        try:
            pc_int = int(pc)
            if 1 <= pc_int <= 50:
                clean_pcs.append(pc_int)
        except:
            pass

    lab_maintenance[lab] = sorted(set(clean_pcs))

    return jsonify({
        "status": "success",
        "message": f"Maintenance saved for Lab {lab}.",
        "lab": lab,
        "maintenance": lab_maintenance[lab]
    })


@app.route('/admin/sitin-report/download/<file_type>')
def download_sitin_report(file_type):
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()
    rows = get_sitin_report_rows(conn)
    conn.close()

    headers = [
        'ID Number',
        'Name',
        'Purpose',
        'Laboratory',
        'Login',
        'Logout',
        'Date'
    ]

    data = []

    for row in rows:
        login_value = row['login_time'] or ''
        logout_value = row['logout_time'] or ''

        login_date = ''
        login_time = ''
        logout_time = '—'

        if login_value:
            try:
                login_dt = datetime.strptime(login_value, '%Y-%m-%d %H:%M:%S')
                login_date = login_dt.strftime('%Y-%m-%d')
                login_time = login_dt.strftime('%H:%M:%S')
            except:
                login_date = login_value[:10]
                login_time = login_value[11:19] if len(login_value) >= 19 else login_value

        if logout_value:
            try:
                logout_dt = datetime.strptime(logout_value, '%Y-%m-%d %H:%M:%S')
                logout_time = logout_dt.strftime('%H:%M:%S')
            except:
                logout_time = logout_value[11:19] if len(logout_value) >= 19 else logout_value

        data.append([
            row['id_number'],
            f"{row['first_name']} {row['last_name']}",
            row['purpose'],
            f"Lab {row['lab']}",
            login_time,
            logout_time,
            login_date
        ])

    if file_type == 'csv':
        output = StringIO()
        writer = csv.writer(output)

        writer.writerow(headers)
        writer.writerows(data)

        return Response(
            output.getvalue(),
            mimetype='text/csv',
            headers={
                'Content-Disposition': 'attachment; filename=ccs_sitin_report.csv'
            }
        )

    elif file_type == 'docx':
        document = Document()

        for section in document.sections:
            section.top_margin = Inches(0.45)
            section.bottom_margin = Inches(0.45)
            section.left_margin = Inches(0.55)
            section.right_margin = Inches(0.55)

        normal_style = document.styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(9)

        top_line = document.add_paragraph()
        top_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

        top_run = top_line.add_run('University of Cebu-Main Campus System Monitoring Reports')
        top_run.font.name = 'Arial'
        top_run.font.size = Pt(8)
        top_run.font.bold = True

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_run = title.add_run('College Of Computer Studies Reports')
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        document.add_paragraph('')

        info_table = document.add_table(rows=1, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_table.autofit = False

        info_table.cell(0, 0).text = datetime.now().strftime('%d/%m/%Y, %H:%M')
        info_table.cell(0, 1).text = f'Total Records: {len(data)}'

        for cell in info_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8)
                    run.font.bold = True

        document.add_paragraph('')

        table = document.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        table.style = 'Table Grid'

        header_cells = table.rows[0].cells

        for index, header in enumerate(headers):
            header_cells[index].text = header
            header_cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            apply_cell_background(header_cells[index], 'F8F8F8')
            apply_cell_padding(header_cells[index], top=90, bottom=90, left=80, right=80)

            paragraph = header_cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        for row_data in data:
            row_cells = table.add_row().cells

            for index, value in enumerate(row_data):
                row_cells[index].text = str(value)
                row_cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                apply_cell_background(row_cells[index], 'FFFFFF')
                apply_cell_padding(row_cells[index], top=80, bottom=80, left=80, right=80)

                paragraph = row_cells[index].paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f"ccs_sitin_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    elif file_type == 'xlsx':
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = 'CCS Sit-in Report'

        sheet.append(headers)

        for row_data in data:
            sheet.append(row_data)

        for column_cells in sheet.columns:
            max_length = 0
            column_letter = column_cells[0].column_letter

            for cell in column_cells:
                value = str(cell.value) if cell.value is not None else ''
                if len(value) > max_length:
                    max_length = len(value)

            sheet.column_dimensions[column_letter].width = min(max_length + 4, 35)

        file_stream = BytesIO()
        workbook.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f"ccs_sitin_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )

    else:
        flash('Invalid download file type.', 'error')
        return redirect(url_for('admin_sitin_report'))


@app.route('/admin/delete-announcement/<int:ann_id>', methods=['POST'])
def delete_announcement(ann_id):
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()
    conn.execute('DELETE FROM announcements WHERE id = ?', (ann_id,))
    conn.commit()
    conn.close()

    flash('Announcement deleted.', 'success')
    return redirect(url_for('admin_dashboard'))


@app.route('/admin/delete-student/<id_number>', methods=['POST'])
def delete_student(id_number):
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()
    conn.execute('DELETE FROM students WHERE id_number = ?', (id_number,))
    conn.commit()
    conn.close()

    flash(f'Student {id_number} deleted successfully.', 'success')
    return redirect(url_for('admin_students'))


@app.route('/admin/logout')
def admin_logout():
    session.clear()
    return redirect(url_for('login'))


@app.route('/admin/get-student/<id_number>')
def get_student(id_number):
    if 'admin_id' not in session:
        return jsonify({"success": False, "message": "Unauthorized"}), 401

    conn = get_db()

    student = conn.execute(
        'SELECT * FROM students WHERE id_number = ?',
        (id_number,)
    ).fetchone()

    if not student:
        conn.close()
        return jsonify({"success": False, "message": "Student not found"})

    session_info = get_student_session_info(conn, id_number)
    conn.close()

    return jsonify({
    "success": True,
    "id_number": student['id_number'],
    "name": f"{student['first_name']} {student['last_name']}",
    "remaining": session_info['remaining_sessions'],
    "completed_sitins": session_info['completed_sitins'],
    "raw_points": session_info['raw_points'],
    "bonus_sessions": session_info['bonus_sessions']
})


@app.route('/admin/search-students')
def search_students():
    if 'admin_id' not in session:
        return jsonify({"success": False, "message": "Unauthorized"}), 401

    query = request.args.get('q', '').strip()

    conn = get_db()

    if not query:
        students = conn.execute('''
            SELECT 
                id_number,
                first_name,
                middle_name,
                last_name,
                course,
                course_level
            FROM students
            ORDER BY last_name, first_name
            LIMIT 50
        ''').fetchall()
    else:
        students = conn.execute('''
            SELECT 
                id_number,
                first_name,
                middle_name,
                last_name,
                course,
                course_level
            FROM students
            WHERE id_number LIKE ? 
               OR first_name LIKE ? 
               OR last_name LIKE ?
            ORDER BY last_name, first_name
            LIMIT 20
        ''', (f'%{query}%', f'%{query}%', f'%{query}%')).fetchall()

    results = []

    for student in students:
        session_info = get_student_session_info(conn, student['id_number'])
        profile_pic_path = os.path.join(app.root_path, 'static', 'uploads', f"{student['id_number']}.png")
        profile_pic = os.path.exists(profile_pic_path)

        results.append({
            'id_number': student['id_number'],
            'first_name': student['first_name'],
            'middle_name': student['middle_name'],
            'last_name': student['last_name'],
            'course': student['course'],
            'course_level': student['course_level'],
            'remaining': session_info['remaining_sessions'],
            'completed_sitins': session_info['completed_sitins'],
            'raw_points': session_info['raw_points'],
            'bonus_sessions': session_info['bonus_sessions'],
            'profile_pic': profile_pic
        })

    conn.close()

    return jsonify({"students": results})


@app.route('/admin/sit-in', methods=['POST'])
def admin_sitin():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    id_number = request.form.get('id_number')
    purpose = request.form.get('purpose')
    lab = request.form.get('lab')
    pc_number = request.form.get('pc_number')
    sitin_date = request.form.get('date')
    sitin_time = request.form.get('time')

    # Auto-fill date/time if empty
    if not sitin_date:
        sitin_date = datetime.now().strftime('%Y-%m-%d')

    if not sitin_time:
        sitin_time = datetime.now().strftime('%H:%M')

    # Required fields, but date/time are already auto-filled above
    if not id_number or not purpose or not lab or not pc_number:
        flash('Please complete all Sit-In fields.', 'error')
        return redirect(url_for('admin_sitin_records'))

    try:
        pc_number_int = int(pc_number)
    except ValueError:
        flash('Invalid PC number selected.', 'error')
        return redirect(url_for('admin_sitin_records'))

    conn = get_db()

    session_info = get_student_session_info(conn, id_number)

    if block_if_pc_maintenance(lab, pc_number_int):
        conn.close()
        flash(f'PC {pc_number_int} in Lab {lab} is currently UNDER MAINTENANCE.', 'error')
        return redirect(url_for('admin_sitin_records'))

    if session_info['remaining_sessions'] <= 0:
        conn.close()
        flash(f'Student {id_number} has no remaining sessions.', 'error')
        return redirect(url_for('admin_sitin_records'))

    existing = conn.execute(
        'SELECT id FROM sitin_records WHERE id_number = ? AND logout_time IS NULL',
        (id_number,)
    ).fetchone()

    if existing:
        conn.close()
        flash(f'Student {id_number} is already sitting in! Log them out first.', 'error')
        return redirect(url_for('admin_sitin_records'))

    active_pc = conn.execute('''
        SELECT id
        FROM sitin_records
        WHERE lab = ?
          AND pc_number = ?
          AND logout_time IS NULL
        LIMIT 1
    ''', (lab, pc_number_int)).fetchone()

    if active_pc:
        conn.close()
        flash(f'PC {pc_number_int} in Lab {lab} is already in use.', 'error')
        return redirect(url_for('admin_sitin_records'))

    login_time = f"{sitin_date} {sitin_time}:00"

    conn.execute('''
        INSERT INTO sitin_records (id_number, purpose, lab, pc_number, login_time)
        VALUES (?, ?, ?, ?, ?)
    ''', (id_number, purpose, lab, pc_number_int, login_time))

    # If the student had a pending reservation that matches this sit-in, mark it as Approved
    try:
        pending_res = conn.execute('''
            SELECT id
            FROM reservations
            WHERE id_number = ?
              AND lab = ?
              AND date = ?
              AND time_in = ?
              AND pc_number = ?
              AND status = 'Pending'
            LIMIT 1
        ''', (id_number, lab, sitin_date, sitin_time, pc_number_int)).fetchone()

        if pending_res:
            conn.execute('''
                UPDATE reservations
                SET status = 'Approved'
                WHERE id = ?
            ''', (pending_res['id'],))

    except Exception:
        pass

    student = conn.execute(
        'SELECT first_name, last_name FROM students WHERE id_number = ?',
        (id_number,)
    ).fetchone()

    conn.commit()
    conn.close()

    if student:
        add_notification(
            id_number,
            f"You have been logged in for sit-in by admin. Purpose: {purpose}, Lab: {lab}, PC {pc_number_int}, Date {sitin_date}, Time {sitin_time}."
        )

    flash(
        f'Student {id_number} successfully sat-in for {sitin_date} {sitin_time} at Lab {lab}, PC {pc_number_int}.',
        'success'
    )

    return redirect(url_for('admin_sitin_records'))

@app.route('/admin/sitin-records')
def admin_sitin_records():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()
    today = date.today().isoformat()

    active_records = conn.execute("""
        SELECT sr.id, sr.id_number, sr.purpose, sr.lab, sr.login_time,
               s.first_name, s.last_name, s.course, s.course_level
        FROM sitin_records sr
        JOIN students s ON sr.id_number = s.id_number
        WHERE sr.logout_time IS NULL
        ORDER BY sr.login_time ASC
    """).fetchall()

    total_sitin = conn.execute(
        "SELECT COUNT(*) as c FROM sitin_records WHERE DATE(login_time) = ?",
        (today,)
    ).fetchone()['c']

    total_logout = conn.execute(
        "SELECT COUNT(*) as c FROM sitin_records WHERE DATE(logout_time) = ?",
        (today,)
    ).fetchone()['c']

    labs = ['524', '526', '528', '530', '542', '544']

    lab_application_rows = conn.execute("""
        SELECT id, lab, application_name
        FROM lab_applications
        ORDER BY lab ASC, application_name ASC
    """).fetchall()

    lab_applications = {lab: [] for lab in labs}

    for row in lab_application_rows:
        lab_key = str(row['lab'])

        if lab_key not in lab_applications:
            lab_applications[lab_key] = []

        lab_applications[lab_key].append({
            'id': row['id'],
            'lab': lab_key,
            'application_name': row['application_name']
        })

    conn.close()

    return render_template(
        'admin_sitin_records.html',
        active_records=active_records,
        total_sitin=total_sitin,
        total_logout=total_logout,
        labs=labs,
        lab_applications=lab_applications,
        admin_user=session['admin_user']
    )


@app.route('/admin/lab-application/add', methods=['POST'])
def admin_add_lab_application():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    lab = request.form.get('lab', '').strip()
    application_name = request.form.get('application_name', '').strip()

    allowed_labs = {'524', '526', '528', '530', '542', '544'}

    if lab not in allowed_labs:
        flash('Invalid laboratory number.', 'error')
        return redirect(url_for('admin_sitin_records'))

    if not application_name:
        flash('Please enter an application name.', 'error')
        return redirect(url_for('admin_sitin_records'))

    conn = get_db()

    try:
        conn.execute("""
            INSERT INTO lab_applications (lab, application_name)
            VALUES (?, ?)
        """, (lab, application_name))

        conn.commit()

        notify_all_students(
            f'Admin added a new available application in Lab {lab}: {application_name}.'
        )

        flash(f'{application_name} added to Lab {lab}.', 'success')

    except sqlite3.IntegrityError:
        flash(f'{application_name} already exists in Lab {lab}.', 'error')

    finally:
        conn.close()

    return redirect(url_for('admin_sitin_records'))


@app.route('/admin/lab-application/delete/<int:application_id>', methods=['POST'])
def admin_delete_lab_application(application_id):
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    application = conn.execute("""
        SELECT lab, application_name
        FROM lab_applications
        WHERE id = ?
    """, (application_id,)).fetchone()

    if not application:
        conn.close()
        flash('Application not found.', 'error')
        return redirect(url_for('admin_sitin_records'))

    conn.execute(
        'DELETE FROM lab_applications WHERE id = ?',
        (application_id,)
    )

    conn.commit()
    conn.close()

    notify_all_students(
        f'Admin removed an application from Lab {application["lab"]}: {application["application_name"]}.'
    )

    flash(
        f'{application["application_name"]} deleted from Lab {application["lab"]}.',
        'success'
    )

    return redirect(url_for('admin_sitin_records'))


@app.route('/admin/sitin-logout/<int:record_id>', methods=['POST'])
def admin_sitin_logout(record_id):
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    cleanliness_checked = request.form.get('evaluation_cleanliness') == '1'
    task_checked = request.form.get('evaluation_task') == '1'
    evaluation_note = request.form.get('evaluation_note', '').strip()

    conn = get_db()

    record_before = conn.execute(
        '''
        SELECT 
            sr.id,
            sr.id_number,
            sr.login_time,
            s.first_name,
            s.last_name
        FROM sitin_records sr
        JOIN students s ON sr.id_number = s.id_number
        WHERE sr.id = ?
          AND sr.logout_time IS NULL
        ''',
        (record_id,)
    ).fetchone()

    if not record_before:
        conn.close()
        flash('Record not found or already logged out.', 'error')
        return redirect(url_for('admin_sitin_records'))

    logout_time = get_local_time()

    duration_row = conn.execute('''
        SELECT CAST(
            (julianday(?) - julianday(?)) * 24 * 60
            AS INTEGER
        ) AS duration_minutes
    ''', (logout_time, record_before['login_time'])).fetchone()

    duration_minutes = duration_row['duration_minutes'] or 0

    # 1 Hour Sit-in is automatic only.
    # Admin cannot manually give this point.
    hours_checked = duration_minutes >= 60

    cleanliness_points = 0.50 if cleanliness_checked else 0
    hours_points = 0.30 if hours_checked else 0
    task_points = 0.20 if task_checked else 0

    evaluation_points = cleanliness_points + hours_points + task_points

    previous_session_info = get_student_session_info(conn, record_before['id_number'])
    previous_bonus_sessions = previous_session_info['bonus_sessions']

    conn.execute('''
        UPDATE sitin_records
        SET logout_time = ?,
            evaluation_cleanliness = ?,
            evaluation_hours = ?,
            evaluation_task = ?,
            evaluation_points = ?,
            evaluation_note = ?
        WHERE id = ?
          AND logout_time IS NULL
    ''', (
        logout_time,
        1 if cleanliness_checked else 0,
        1 if hours_checked else 0,
        1 if task_checked else 0,
        evaluation_points,
        evaluation_note,
        record_id
    ))

    conn.commit()

    session_info = get_student_session_info(conn, record_before['id_number'])
    new_bonus_sessions = session_info['bonus_sessions']
    extra_sessions_added = new_bonus_sessions - previous_bonus_sessions

    conn.close()

    message = (
        f"You have been logged out from sit-in by admin. "
        f"You received {evaluation_points:.2f} raw point(s)."
    )

    if cleanliness_checked:
        message += " Cleanliness: +0.50."

    if hours_checked:
        message += " 1 hour sit-in reached: +0.30."
    else:
        message += " 1 hour sit-in was not reached."

    if task_checked:
        message += " Task completed: +0.20."

    if evaluation_points == 0:
        message += " No evaluation criteria were achieved."

    if extra_sessions_added > 0:
        message += (
            f" Congratulations! Your raw points reached a new multiple of 3. "
            f"You earned {extra_sessions_added} extra session."
        )

    if evaluation_note:
        message += f" Admin note: {evaluation_note}"

    add_notification(record_before['id_number'], message)

    flash_message = (
        f'{record_before["first_name"]} {record_before["last_name"]} '
        f'({record_before["id_number"]}) has been logged out. '
        f'Evaluation points given: {evaluation_points:.2f}. '
        f'Total raw points: {session_info["raw_points"]:.2f}. '
        f'Total extra sessions earned: {session_info["bonus_sessions"]}. '
        f'Remaining sessions: {session_info["remaining_sessions"]}.'
    )

    if hours_checked:
        flash_message += ' Student reached at least 1 hour, so +0.30 was added automatically.'
    else:
        flash_message += ' Student did not reach 1 hour, so +0.30 was not added.'

    if extra_sessions_added > 0:
        flash_message += f' New extra session added: +{extra_sessions_added}.'

    flash(flash_message, 'success')

    return redirect(url_for('admin_sitin_records'))


@app.route('/admin/reservations')
def admin_reservations():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    reservations = conn.execute('''
        SELECT 
            r.*,
            s.first_name,
            s.last_name,
            s.course,
            s.course_level
        FROM reservations r
        JOIN students s ON r.id_number = s.id_number
        ORDER BY r.created_at DESC
    ''').fetchall()

    reservations_enabled = is_reservation_enabled(conn)
    
    conn.close()

    return render_template(
        'admin_reservations.html',
        reservations=reservations,
        reservations_enabled=reservations_enabled,
        admin_user=session['admin_user']
    )


@app.route('/admin/reservation/toggle', methods=['POST'])
def toggle_reservation():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()
    is_now_enabled = toggle_reservation_status(conn)
    conn.close()

    status_text = "enabled" if is_now_enabled else "disabled"
    flash(f'Reservations have been {status_text}.', 'success')
    return redirect(url_for('admin_reservations'))


@app.route('/admin/reservation/approve/<int:res_id>', methods=['POST'])
def approve_reservation(res_id):
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    reservation = conn.execute("""
        SELECT *
        FROM reservations
        WHERE id = ?
    """, (res_id,)).fetchone()

    if not reservation:
        conn.close()
        flash('Reservation not found.', 'error')
        return redirect(url_for('admin_reservations'))

    session_info = get_student_session_info(conn, reservation['id_number'])

    if session_info['remaining_sessions'] <= 0:
        conn.close()
        flash('Cannot approve. Student has no remaining sessions.', 'error')
        return redirect(url_for('admin_reservations'))

    conn.execute(
        "UPDATE reservations SET status = 'Approved' WHERE id = ?",
        (res_id,)
    )

    conn.commit()
    conn.close()

    add_notification(
        reservation['id_number'],
        f"Your reservation has been approved by admin. Purpose: {reservation['purpose']}, Lab: {reservation['lab']}, PC: {reservation['pc_number']}, Date: {reservation['date']}, Time: {reservation['time_in']}."
    )

    flash('Reservation approved.', 'success')
    return redirect(url_for('admin_reservations'))


@app.route('/admin/reservation/reject/<int:res_id>', methods=['POST'])
def reject_reservation(res_id):
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    reservation = conn.execute("""
        SELECT *
        FROM reservations
        WHERE id = ?
    """, (res_id,)).fetchone()

    conn.execute(
        "UPDATE reservations SET status = 'Rejected' WHERE id = ?",
        (res_id,)
    )

    conn.commit()
    conn.close()

    if reservation:
        add_notification(
            reservation['id_number'],
            f"Your reservation for Lab {reservation['lab']} on {reservation['date']} at {reservation['time_in']} has been rejected by admin."
        )

    flash('Reservation rejected.', 'success')
    return redirect(url_for('admin_reservations'))


@app.route('/admin/students')
def admin_students():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    students = conn.execute(
        'SELECT * FROM students ORDER BY last_name, first_name'
    ).fetchall()

    total_students = len(students)

    students_with_sessions = []

    for student in students:
        session_info = get_student_session_info(conn, student['id_number'])

        student_dict = dict(student)
        student_dict['remaining_sessions'] = session_info['remaining_sessions']
        student_dict['completed_sitins'] = session_info['completed_sitins']
        student_dict['bonus_sessions'] = session_info['bonus_sessions']
        student_dict['raw_points'] = session_info['raw_points']

        students_with_sessions.append(student_dict)

    leaderboard_rows = conn.execute("""
        SELECT
            s.id_number,
            s.first_name,
            s.middle_name,
            s.last_name,
            s.course,
            s.course_level,

            COUNT(sr.id) AS total_sitins,

            SUM(
                CASE
                    WHEN sr.logout_time IS NOT NULL THEN 1
                    ELSE 0
                END
            ) AS completed_sitins,

            SUM(
                CASE
                    WHEN sr.logout_time IS NULL AND sr.id IS NOT NULL THEN 1
                    ELSE 0
                END
            ) AS active_sitins,

            COALESCE(
                SUM(
                    CASE
                        WHEN sr.logout_time IS NOT NULL THEN sr.evaluation_points
                        ELSE 0
                    END
                ),
                0
            ) AS raw_points,

            CAST(
                COALESCE(
                    SUM(
                        CASE
                            WHEN sr.logout_time IS NOT NULL THEN sr.evaluation_points
                            ELSE 0
                        END
                    ),
                    0
                ) / 3 AS INTEGER
            ) AS bonus_sessions

        FROM students s
        LEFT JOIN sitin_records sr
            ON s.id_number = sr.id_number

        GROUP BY
            s.id_number,
            s.first_name,
            s.middle_name,
            s.last_name,
            s.course,
            s.course_level

        ORDER BY
            raw_points DESC,
            completed_sitins DESC,
            s.last_name ASC,
            s.first_name ASC
    """).fetchall()

    leaderboard_students = []

    for row in leaderboard_rows:
        row_dict = dict(row)

        row_dict['total_sitins'] = row_dict['total_sitins'] or 0
        row_dict['completed_sitins'] = row_dict['completed_sitins'] or 0
        row_dict['active_sitins'] = row_dict['active_sitins'] or 0
        row_dict['raw_points'] = float(row_dict['raw_points'] or 0)
        row_dict['bonus_sessions'] = int(row_dict['bonus_sessions'] or 0)

        leaderboard_students.append(row_dict)

    conn.close()

    return render_template(
        'admin_students.html',
        students=students_with_sessions,
        total_students=total_students,
        leaderboard_students=leaderboard_students,
        admin_user=session['admin_user']
    )


@app.route('/admin/reset-sessions', methods=['POST'])
def admin_reset_sessions():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    # Step 1: Insert a new session reset timestamp
    conn.execute("INSERT INTO session_resets (reset_at) VALUES (CURRENT_TIMESTAMP)")

    # Step 2: Reset all evaluation points of past sit-ins
    conn.execute("""
        UPDATE sitin_records
        SET evaluation_cleanliness = 0,
            evaluation_hours = 0,
            evaluation_task = 0,
            evaluation_points = 0,
            evaluation_note = NULL
        WHERE logout_time IS NOT NULL
    """)

    conn.commit()
    conn.close()

    # Step 3: Notify all students
    notify_all_students(
        'Admin reset all student sessions and raw evaluation points back to zero.'
    )

    # Step 4: Flash success message to admin
    flash('All student sessions have been reset to 30 and raw evaluation points cleared.', 'success')

    return redirect(url_for('admin_students'))

@app.route('/admin/sitin-reports')
def admin_sitin_reports():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    total_sitin = conn.execute(
        'SELECT COUNT(*) as c FROM sitin_records'
    ).fetchone()['c']

    total_completed = conn.execute(
        'SELECT COUNT(*) as c FROM sitin_records WHERE logout_time IS NOT NULL'
    ).fetchone()['c']

    purpose_stats = conn.execute('''
        SELECT purpose, COUNT(*) as count 
        FROM sitin_records 
        GROUP BY purpose 
        ORDER BY count DESC
    ''').fetchall()

    lab_stats = conn.execute('''
        SELECT lab, COUNT(*) as count 
        FROM sitin_records 
        GROUP BY lab 
        ORDER BY count DESC
    ''').fetchall()

    conn.close()

    return render_template(
        'admin_sitin_reports.html',
        total_sitin=total_sitin,
        total_completed=total_completed,
        purpose_stats=purpose_stats,
        lab_stats=lab_stats,
        admin_user=session['admin_user']
    )


@app.route('/admin/feedback-reports')
def admin_feedback_reports():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    feedbacks = conn.execute('''
        SELECT sr.id, sr.id_number, sr.purpose, sr.lab, sr.login_time,
               sr.logout_time, sr.rating, sr.feedback,
               s.first_name, s.last_name, s.course
        FROM sitin_records sr
        JOIN students s ON sr.id_number = s.id_number
        WHERE sr.feedback IS NOT NULL
        ORDER BY sr.login_time DESC
    ''').fetchall()

    avg_rating = conn.execute('''
        SELECT AVG(rating) as avg_rating 
        FROM sitin_records 
        WHERE rating IS NOT NULL
    ''').fetchone()['avg_rating']

    total_feedbacks = len(feedbacks)

    conn.close()

    return render_template(
        'admin_feedback_reports.html',
        feedbacks=feedbacks,
        avg_rating=avg_rating,
        total_feedbacks=total_feedbacks,
        admin_user=session['admin_user']
    )


@app.route('/history')
def history():
    if 'student_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    records = conn.execute('''
        SELECT *
        FROM sitin_records
        WHERE id_number = ?
        ORDER BY login_time DESC
    ''', (session['student_id'],)).fetchall()

    conn.close()

    notifications, notification_count = get_student_notifications(session['student_id'])

    return render_template(
        'history.html',
        records=records,
        notifications=notifications,
        notification_count=notification_count
    )


@app.route('/submit-feedback', methods=['POST'])
def submit_feedback():
    if 'student_id' not in session:
        return redirect(url_for('login'))

    record_id = request.form.get('record_id')
    rating = request.form.get('rating')
    feedback = request.form.get('feedback', '').strip()

    if not record_id or not rating or not feedback:
        flash('Please complete the feedback form.', 'error')
        return redirect(url_for('history'))

    conn = get_db()

    record = conn.execute('''
        SELECT * FROM sitin_records
        WHERE id = ? AND id_number = ?
    ''', (record_id, session['student_id'])).fetchone()

    if not record:
        conn.close()
        flash('Invalid sit-in record.', 'error')
        return redirect(url_for('history'))

    conn.execute('''
        UPDATE sitin_records
        SET rating = ?, feedback = ?
        WHERE id = ?
    ''', (rating, feedback, record_id))

    conn.commit()
    conn.close()

    add_notification(
        session['student_id'],
        'Your feedback was submitted successfully. Thank you for your response.'
    )

    flash('Feedback submitted successfully!', 'success')
    return redirect(url_for('history'))


@app.route('/admin/view-sitin-records')
def view_sitin_records():
    if 'admin_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    search_id = request.args.get('search_id', '').strip()
    page = request.args.get('page', 1, type=int)
    per_page = 10

    if page < 1:
        page = 1

    offset = (page - 1) * per_page

    base_query = '''
        FROM sitin_records sr
        JOIN students s ON sr.id_number = s.id_number
    '''

    params = []

    if search_id:
        base_query += '''
            WHERE sr.id_number LIKE ?
               OR s.first_name LIKE ?
               OR s.last_name LIKE ?
               OR sr.purpose LIKE ?
               OR sr.lab LIKE ?
        '''
        search_value = f'%{search_id}%'
        params.extend([
            search_value,
            search_value,
            search_value,
            search_value,
            search_value
        ])

    total_records = conn.execute(
        f'SELECT COUNT(*) AS count {base_query}',
        params
    ).fetchone()['count']

    total_pages = (total_records + per_page - 1) // per_page

    records = conn.execute(
        f'''
        SELECT 
            sr.id,
            sr.id_number,
            sr.purpose,
            sr.lab,
            sr.login_time,
            sr.logout_time,
            s.first_name,
            s.last_name,
            s.course
        {base_query}
        ORDER BY sr.login_time DESC
        LIMIT ? OFFSET ?
        ''',
        params + [per_page, offset]
    ).fetchall()

    conn.close()

    return render_template(
        'admin_view_sitin_records.html',
        records=records,
        admin_user=session['admin_user'],
        search_id=search_id,
        page=page,
        per_page=per_page,
        total_pages=total_pages,
        total_records=total_records
    )


@app.route('/notification/read/<int:notif_id>')
def read_notification(notif_id):
    if 'student_id' not in session:
        return redirect(url_for('login'))

    conn = get_db()

    conn.execute('''
        UPDATE notifications
        SET is_read = 1
        WHERE id = ?
          AND id_number = ?
    ''', (notif_id, session['student_id']))

    conn.commit()
    conn.close()

    return redirect(request.referrer or url_for('dashboard'))

@app.route('/notifications/mark-read', methods=['POST'])
def mark_notifications_read():
    if 'student_id' not in session:
        return jsonify({'success': False, 'message': 'Unauthorized'}), 401

    conn = get_db()

    conn.execute('''
        UPDATE notifications
        SET is_read = 1
        WHERE id_number = ?
          AND is_read = 0
    ''', (session['student_id'],))

    conn.commit()
    conn.close()

    return jsonify({'success': True})


if __name__ == '__main__':
    init_db()
    app.run(host='0.0.0.0', port=1234, debug=True)